# import os
# import pandas as pd
# import requests
# from datetime import datetime
# from docx import Document
# from docx.shared import Pt
# import mimetypes
# from email.message import EmailMessage

# OLLAMA_URL = "http://localhost:11434"

# # ------------------------------------
# # ✅ Ollama Call
# # ------------------------------------
# def call_ollama_model(prompt, model="llama3", timeout=300):
#     payload = {
#         "model": model,
#         "prompt": prompt,
#         "stream": False
#     }
#     try:
#         response = requests.post(f"{OLLAMA_URL}/api/generate", json=payload, timeout=timeout)
#         response.raise_for_status()
#         return response.json().get("response", "").strip()
#     except requests.exceptions.Timeout:
#         return "❌ Ollama timed out."
#     except Exception as e:
#         return f"❌ Ollama Error: {e}"

# # ------------------------------------
# # ✅ Helpers
# # ------------------------------------
# def extract_attendance_details(attendance_files):
#     if not attendance_files:
#         return "N/A", 0

#     attendance_details = []
#     for file in attendance_files:
#         df = pd.read_csv(file)
#         df['Event Date'] = pd.to_datetime(df['UTC Event Timestamp'], errors='coerce').dt.date
#         most_active_date = df['Event Date'].value_counts().idxmax()
#         filtered_df = df[df['Event Date'] == most_active_date]
#         unique_attendance = filtered_df['Participant Id'].nunique()
#         attendance_details.append((most_active_date, unique_attendance))

#     total_attendance = sum([c for _, c in attendance_details])
#     earliest_date = min([d for d, _ in attendance_details])
#     return str(earliest_date), total_attendance

# def parse_vtt_files(vtt_files):
#     transcript = []
#     for file in vtt_files:
#         with open(file, 'r') as f:
#             lines = f.readlines()
#         for line in lines:
#             if '-->' not in line and line.strip():
#                 transcript.append(line.strip())
#     return ' '.join(transcript)

# def polish_question(question, model="llama3"):
#     prompt = f"""
#     Improve the clarity and grammar of the question below.
#     Return ONLY the improved question text, no extra lines.

#     Question:
#     \"\"\"{question}\"\"\"
#     """
#     return call_ollama_model(prompt, model)

# def polish_answer(question, answer, model="llama3"):
#     prompt = f"""
#     Improve the grammar and clarity of the answer text below.
#     Keep it short, clear, and professional.
#     Return ONLY the improved answer, no extra lines.

#     Answer:
#     \"\"\"{answer}\"\"\"
#     """
#     return call_ollama_model(prompt, model)

# def clean_ai_response(text):
#     return text.strip().replace("\n", " ")

# def extract_answer_from_transcript(question, transcript, model="llama3"):
#     prompt = f"""
#     Use this transcript to find an answer to the question ONLY if it was verbally answered.
#     If not, say: "Answer could not be deciphered from the transcript."

#     Transcript:
#     {transcript}

#     Question:
#     {question}
#     """
#     return call_ollama_model(prompt, model)

# def extract_feedback(feedback_files, selected_title):
#     combined_df = pd.concat([pd.read_excel(f) for f in feedback_files], ignore_index=True)
#     filtered_df = combined_df[combined_df['Training Title'] == selected_title]

#     if filtered_df.empty:
#         raise ValueError(f"No feedback for: {selected_title}")

#     avg_rating = round(filtered_df["User's Course Rating"].mean(), 2)
#     blocks = []
#     for _, row in filtered_df.iterrows():
#         title = str(row.get('User Course Review Title', '')).strip()
#         desc = str(row.get('User Course Review Desc', '')).strip()
#         if title or desc:
#             blocks.append(f"Title: {title}\nDescription: {desc}")
#     return avg_rating, "\n\n".join(blocks)

# def generate_ai_summary(session_title, avg_rating, feedback_text, model="llama3"):
#     prompt = f"""
#     Write a short professional summary of the session below using the feedback.
#     If feedback is empty, use the rating: {avg_rating}.

#     Feedback:
#     {feedback_text}
#     """
#     return call_ollama_model(prompt, model)

# def create_email_body(session_title, session_date, attendance_count, avg_rating, ai_summary, repo_link, sp_link, open_questions):
#     return f"""
# Hi All,

# Thank you for attending {session_title} on {session_date}.

# The session received an average rating of {avg_rating}/5 with {attendance_count} attendees.
# Key points from feedback:
# {ai_summary}

# Q&A report is here: {repo_link}
# Open questions: {open_questions}
# Links are updated on SharePoint: {sp_link}

# Best regards
# """

# def create_eml_with_attachment(subject, body, attachment_paths, output_path="Session_Email.eml"):
#     msg = EmailMessage()
#     msg['Subject'] = subject
#     msg['From'] = "sender@example.com"
#     msg['To'] = "recipient@example.com"
#     msg.set_content(body)

#     for path in attachment_paths:
#         mime_type, _ = mimetypes.guess_type(path)
#         if mime_type is None:
#             mime_type = 'application/octet-stream'
#         main_type, sub_type = mime_type.split('/', 1)
#         with open(path, 'rb') as f:
#             msg.add_attachment(f.read(), maintype=main_type, subtype=sub_type, filename=os.path.basename(path))

#     with open(output_path, 'wb') as f:
#         f.write(bytes(msg))

#     return output_path

# def filter_irrelevant_content(content, model="llama3"):
#     prompt = f"""
#     Determine if this is a valid question.
#     If valid, return it.
#     If not, return empty.

#     Content:
#     \"\"\"{content}\"\"\"
#     """
#     result = call_ollama_model(prompt, model).strip()
#     return result if result else None

# def generate_qa_report_with_transcript(qna_files, transcript_files=None, model="llama3"):
#     dfs = [pd.read_csv(f) for f in qna_files]
#     df = pd.concat(dfs, ignore_index=True)
#     df['Content'] = df['Content'].fillna('')
#     df = df.drop_duplicates(subset=['Conversation Id', 'Content'])

#     questions = df[df['Type'].isin(['Question', 'Announcement'])]
#     responses = df[df['Type'] == 'Response']
#     response_dict = responses.groupby('Conversation Id')['Content'].apply(list).to_dict()

#     transcript = parse_vtt_files(transcript_files) if transcript_files else ""

#     qa_pairs = []
#     open_questions = []
#     for _, row in questions.iterrows():
#         conv_id = row['Conversation Id']
#         question_text = row['Content'].strip()
#         answers = response_dict.get(conv_id, [])
#         answer_text = '\n'.join([a for a in answers if a.strip()])

#         if not answer_text:
#             open_questions.append(question_text)
#         else:
#             qa_pairs.append((question_text, answer_text))

#     # Q&A doc
#     doc = Document()
#     doc.add_heading('Session Q&A', level=0)

#     for q, a in qa_pairs:
#         polished_q = clean_ai_response(polish_question(q))
#         full_answer = a
#         if "answered verbally" in a.lower() and transcript:
#             trans_answer = extract_answer_from_transcript(q, transcript)
#             full_answer += f"\n\n{trans_answer}"
#         polished_a = clean_ai_response(polish_answer(polished_q, full_answer))
#         doc.add_heading(f"Q: {polished_q}", level=2)
#         doc.add_paragraph(f"A: {polished_a}")
#     qa_path = "QA_Report.docx"
#     doc.save(qa_path)

#     # Open questions doc
#     open_doc = Document()
#     open_doc.add_heading("Open Questions", level=0)
#     final_open = []
#     for oq in open_questions:
#         res = filter_irrelevant_content(oq)
#         if res:
#             final_open.append(res)
#             open_doc.add_heading(f"Q: {res}", level=2)
#     open_path = "Open_Questions.docx"
#     open_doc.save(open_path)

#     return qa_path, open_path, len(final_open)

# # ------------------------------------
# # ✅ FINAL WRAPPER
# # ------------------------------------
# def process_training_files(attendance_files, qna_files, transcript_files, feedback_files, training_title_input, repository_link, sharepoint_link):
#     def save_uploaded(files, subfolder):
#         paths = []
#         folder = os.path.join("uploads", subfolder)
#         os.makedirs(folder, exist_ok=True)
#         for f in files:
#             path = os.path.join(folder, f.filename)
#             f.save(path)
#             paths.append(path)
#         return paths

#     att_paths = save_uploaded(attendance_files, "attendance")
#     qna_paths = save_uploaded(qna_files, "qna")
#     tr_paths = save_uploaded(transcript_files, "transcript")
#     fb_paths = save_uploaded(feedback_files, "feedback")

#     feedback_dfs = [pd.read_excel(f) for f in fb_paths]
#     feedback_df = pd.concat(feedback_dfs, ignore_index=True)
#     unique_titles = feedback_df['Training Title'].dropna().unique() if 'Training Title' in feedback_df.columns else []
#     if training_title_input not in unique_titles:
#         raise ValueError("Training title does not match any feedback data.")

#     session_date, attendance_count = extract_attendance_details(att_paths)
#     avg_rating, feedback_text = extract_feedback(fb_paths, training_title_input)
#     qa_docx, open_docx, open_qs = generate_qa_report_with_transcript(qna_paths, tr_paths)
#     summary = generate_ai_summary(training_title_input, avg_rating, feedback_text)
#     email_body = create_email_body(training_title_input, session_date, attendance_count, avg_rating, summary, repository_link, sharepoint_link, open_qs)

#     attachments = [qa_docx, open_docx] + tr_paths
#     eml = create_eml_with_attachment(f"Summary: {training_title_input} on {session_date}", email_body, attachments)

#     return {"eml_path": eml}

# import os
# import pandas as pd
# import requests
# from datetime import datetime
# from docx import Document
# from docx.shared import Pt
# import mimetypes
# from email.message import EmailMessage

# OLLAMA_URL = "http://localhost:11434"

# # ------------------------------------
# # ✅ Ollama Call
# # ------------------------------------
# def call_ollama_model(prompt, model="llama3", timeout=300):
#     payload = {
#         "model": model,
#         "prompt": prompt,
#         "stream": False
#     }
#     try:
#         response = requests.post(f"{OLLAMA_URL}/api/generate", json=payload, timeout=timeout)
#         response.raise_for_status()
#         return response.json().get("response", "").strip()
#     except requests.exceptions.Timeout:
#         return "❌ Ollama timed out."
#     except Exception as e:
#         return f"❌ Ollama Error: {e}"


# # ------------------------------------
# # ✅ Helpers
# # ------------------------------------
# def extract_attendance_details(attendance_files):
#     if not attendance_files:
#         return "N/A", 0

#     attendance_details = []
#     for file in attendance_files:
#         df = pd.read_csv(file)
#         df['Event Date'] = pd.to_datetime(df['UTC Event Timestamp'], errors='coerce').dt.date
#         most_active_date = df['Event Date'].value_counts().idxmax()
#         filtered_df = df[df['Event Date'] == most_active_date]
#         unique_attendance = filtered_df['Participant Id'].nunique()
#         attendance_details.append((most_active_date, unique_attendance))

#     total_attendance = sum([count for _, count in attendance_details])
#     earliest_date = min([date for date, _ in attendance_details])
#     return str(earliest_date), total_attendance


# def parse_vtt_files(vtt_files):
#     transcript = []
#     for file in vtt_files:
#         with open(file, 'r') as f:
#             lines = f.readlines()
#         for line in lines:
#             if '-->' not in line and line.strip():
#                 transcript.append(line.strip())
#     return ' '.join(transcript)


# def polish_question(question, model="llama3"):
#     prompt = f"""
#     Improve the clarity and grammar of the question below.
#     Return ONLY the improved question text, no extra lines.

#     Question:
#     \"\"\"{question}\"\"\"
#     """
#     return call_ollama_model(prompt, model)


# def polish_answer(question, answer, model="llama3"):
#     prompt = f"""
#     Improve the grammar and clarity of the answer text below.
#     Keep it short, clear, and professional.
#     Return ONLY the improved answer, no extra lines.

#     Answer:
#     \"\"\"{answer}\"\"\"
#     """
#     return call_ollama_model(prompt, model)


# def clean_ai_response(text):
#     return text.strip().replace("\n", " ")


# def extract_answer_from_transcript(question, transcript, model="llama3"):
#     prompt = f"""
#     Use this transcript to find an answer to the question ONLY if it was verbally answered.
#     If not, say: "Answer could not be deciphered from the transcript."

#     Transcript:
#     {transcript}

#     Question:
#     {question}
#     """
#     return call_ollama_model(prompt, model)


# def extract_feedback(feedback_files, selected_title):
#     combined_df = pd.concat([pd.read_excel(f) for f in feedback_files], ignore_index=True)
#     filtered_df = combined_df[combined_df['Training Title'] == selected_title]

#     if filtered_df.empty:
#         raise ValueError(f"No feedback for: {selected_title}")

#     avg_rating = round(filtered_df["User's Course Rating"].mean(), 2)
#     blocks = []
#     for _, row in filtered_df.iterrows():
#         title = str(row.get('User Course Review Title', '')).strip()
#         desc = str(row.get('User Course Review Desc', '')).strip()
#         if title or desc:
#             blocks.append(f"Title: {title}\nDescription: {desc}")
#     return avg_rating, "\n\n".join(blocks)


# def generate_ai_summary(session_title, avg_rating, feedback_text, model="llama3"):
#     prompt = f"""
#     Write a short professional summary of the session below using the feedback.
#     If feedback is empty, use the rating: {avg_rating}.

#     Feedback:
#     {feedback_text}
#     """
#     return call_ollama_model(prompt, model)


# def create_email_body(session_title, session_date, attendance_count, avg_rating, ai_summary, repo_link, sp_link, open_questions):
#     return f"""
# Hi All,

# Thank you for attending {session_title} on {session_date}.

# The session received an average rating of {avg_rating}/5 with {attendance_count} attendees.
# Key points from feedback:
# {ai_summary}

# Q&A report is here: {repo_link}
# Open questions: {open_questions}
# Links are updated on SharePoint: {sp_link}

# Best regards
# """


# def create_eml_with_attachment(subject, body, attachment_paths, output_path="Session_Email.eml"):
#     msg = EmailMessage()
#     msg['Subject'] = subject
#     msg['From'] = "sender@example.com"
#     msg['To'] = "recipient@example.com"
#     msg.set_content(body)

#     for path in attachment_paths:
#         mime_type, _ = mimetypes.guess_type(path)
#         if mime_type is None:
#             mime_type = 'application/octet-stream'
#         main_type, sub_type = mime_type.split('/', 1)
#         with open(path, 'rb') as f:
#             msg.add_attachment(f.read(), maintype=main_type, subtype=sub_type, filename=os.path.basename(path))

#     with open(output_path, 'wb') as f:
#         f.write(bytes(msg))

#     return output_path


# def filter_irrelevant_content(content, model="llama3"):
#     prompt = f"""
#     Determine if this is a valid question.
#     If valid, return it.
#     If not, return empty.

#     Content:
#     \"\"\"{content}\"\"\"
#     """
#     result = call_ollama_model(prompt, model).strip()
#     return result if result else None


# def generate_qa_report_with_transcript(qna_files, transcript_files=None, model="llama3"):
#     dfs = [pd.read_csv(f) for f in qna_files]
#     df = pd.concat(dfs, ignore_index=True)
#     df['Content'] = df['Content'].fillna('')
#     df = df.drop_duplicates(subset=['Conversation Id', 'Content'])

#     questions = df[df['Type'].isin(['Question', 'Announcement'])]
#     responses = df[df['Type'] == 'Response']
#     response_dict = responses.groupby('Conversation Id')['Content'].apply(list).to_dict()

#     transcript = parse_vtt_files(transcript_files) if transcript_files else ""

#     qa_pairs = []
#     open_questions = []

#     for _, row in questions.iterrows():
#         conv_id = row['Conversation Id']
#         question_text = row['Content'].strip()
#         answers = response_dict.get(conv_id, [])
#         answer_text = '\n'.join([a for a in answers if a.strip()])

#         if not answer_text:
#             open_questions.append(question_text)
#         else:
#             qa_pairs.append((question_text, answer_text))

#     # Q&A Document
#     doc = Document()
#     doc.add_heading('Session Q&A', level=0)

#     for q, a in qa_pairs:
#         polished_q = clean_ai_response(polish_question(q))
#         full_answer = a
#         if "answered verbally" in a.lower() and transcript:
#             trans_answer = extract_answer_from_transcript(q, transcript)
#             full_answer += f"\n\n{trans_answer}"
#         polished_a = clean_ai_response(polish_answer(polished_q, full_answer))
#         doc.add_heading(f"Q: {polished_q}", level=2)
#         doc.add_paragraph(f"A: {polished_a}")

#     qa_path = "QA_Report.docx"
#     doc.save(qa_path)

#     # Open Questions Document
#     open_doc = Document()
#     open_doc.add_heading("Open Questions", level=0)

#     final_open = []
#     for oq in open_questions:
#         res = filter_irrelevant_content(oq)
#         if res:
#             final_open.append(res)
#             open_doc.add_heading(f"Q: {res}", level=2)

#     open_path = "Open_Questions.docx"
#     open_doc.save(open_path)

#     return qa_path, open_path, len(final_open)


# # ------------------------------------
# # ✅ FINAL WRAPPER
# # ------------------------------------
# def process_training_files(
#     attendance_files,
#     qna_files,
#     transcript_files,
#     feedback_files,
#     training_title_input,
#     repository_link,
#     sharepoint_link
# ):
#     def save_uploaded(files, subfolder):
#         paths = []
#         folder = os.path.join("uploads", subfolder)
#         os.makedirs(folder, exist_ok=True)
#         for f in files:
#             path = os.path.join(folder, f.filename)
#             f.save(path)
#             paths.append(path)
#         return paths

#     att_paths = save_uploaded(attendance_files, "attendance")
#     qna_paths = save_uploaded(qna_files, "qna")
#     tr_paths = save_uploaded(transcript_files, "transcript")
#     fb_paths = save_uploaded(feedback_files, "feedback")

#     feedback_dfs = [pd.read_excel(f) for f in fb_paths]
#     feedback_df = pd.concat(feedback_dfs, ignore_index=True)
#     unique_titles = feedback_df['Training Title'].dropna().unique() if 'Training Title' in feedback_df.columns else []

#     if training_title_input not in unique_titles:
#         raise ValueError("Training title does not match any feedback data.")

#     session_date, attendance_count = extract_attendance_details(att_paths)
#     avg_rating, feedback_text = extract_feedback(fb_paths, training_title_input)
#     qa_docx, open_docx, open_qs = generate_qa_report_with_transcript(qna_paths, tr_paths)
#     summary = generate_ai_summary(training_title_input, avg_rating, feedback_text)
#     email_body = create_email_body(
#         training_title_input,
#         session_date,
#         attendance_count,
#         avg_rating,
#         summary,
#         repository_link,
#         sharepoint_link,
#         open_qs
#     )

#     attachments = [qa_docx, open_docx] + tr_paths
#     eml = create_eml_with_attachment(
#         f"Summary: {training_title_input} on {session_date}",
#         email_body,
#         attachments
#     )

#     return {"eml_path": eml}



import os
import uuid
import pandas as pd
from datetime import datetime
from docx import Document
from docx.shared import Pt
from flask import jsonify
import requests
from email.message import EmailMessage
import time
import mimetypes

OLLAMA_URL = "http://localhost:11434"


# --- LLM CALL ---
def call_ollama_model(prompt, model="llama3", timeout=300):
    payload = {
        "model": model,
        "prompt": prompt,
        "stream": False
    }

    try:
        response = requests.post(f"{OLLAMA_URL}/api/generate", json=payload)
        response.raise_for_status()
        return response.json().get("response", "").strip()
    except requests.exceptions.Timeout:
        return "Error generating response due to timeout."
    except requests.exceptions.RequestException:
        return "Error generating response due to network or API error."
    except Exception:
        return "Error generating response due to an unexpected error."


# --- TRAINING TITLE VALIDATION ---
def validate_training_title(feedback_paths, title_to_check):
    feedback_dfs = [pd.read_excel(f) for f in feedback_paths]
    feedback_df = pd.concat(feedback_dfs, ignore_index=True)

    if 'Training Title' not in feedback_df.columns:
        raise ValueError("Missing 'Training Title' column in feedback.")

    valid_titles = feedback_df['Training Title'].dropna().unique()
    if title_to_check not in valid_titles:
        raise ValueError(f"Training title '{title_to_check}' not found in feedback.")

    return feedback_df


# --- ATTENDANCE FUNCTION ---
def extract_attendance_details(attendance_files):
    if not attendance_files:
        return "N/A", 0

    attendance_details = []
    for file in attendance_files:
        try:
            df = pd.read_csv(file)
            df['Event Date'] = pd.to_datetime(df['UTC Event Timestamp'], errors='coerce').dt.date
            most_active_date = df['Event Date'].value_counts().idxmax()
            filtered_df = df[df['Event Date'] == most_active_date]
            unique_attendance = filtered_df['Participant Id'].nunique()
            attendance_details.append((most_active_date, unique_attendance))
        except Exception:
            pass

    if not attendance_details:
        return "N/A", 0

    total_attendance = sum([count for _, count in attendance_details])
    earliest_date = min([date for date, _ in attendance_details])

    return str(earliest_date), total_attendance


# --- TRANSCRIPT PARSING ---
def parse_vtt_files(vtt_files):
    full_transcript = []
    for vtt_file in vtt_files:
        try:
            with open(vtt_file, 'r') as file:
                lines = file.readlines()
            for line in lines:
                if '-->' not in line and line.strip():
                    full_transcript.append(line.strip())
        except Exception:
            pass
    return ' '.join(full_transcript)


# --- TRANSCRIPT Q&A EXTRACTION ---
def extract_answer_from_transcript(question, transcript, model="llama3"):
    prompt = f"""
    Use the transcript to find the answers to the question only/strictly where it is mentioned/written (in the answer section) 
    that they were answered verbally during the meeting.
    If no verbal answer is found, return saying "Answer could not be deciphered from the transcript".

    Given the following meeting transcript, answer the question:

    Transcript: {transcript}
    Question: {question}
    """
    return call_ollama_model(prompt, model)


# --- POLISH QUESTION ---
def polish_question(question, model="llama3"):
    prompt = f""""
    You are a helpful assistant. 

    Improve the clarity and grammar of the following question without changing its meaning.
    No Need to Mention/Show/Inform the user that it has been polished/refined/cleaned by you. 
    Just return the text in a polished language.

    ⚠️ Return **ONLY** the improved question.

    Question:
    \"\"\"{question}\"\"\"
    """
    return call_ollama_model(prompt, model)


# --- POLISH ANSWER ---
def polish_answer(question, answer, model="llama3"):
    prompt = f"""
    You are a helpful assistant. 

    Improve the grammar, structure, and clarity of the following answer.
    Keep technical accuracy intact. No introductions, comments, or third-person voice.

    ⚠️ Return **ONLY** the improved answer.

    Answer:
    \"\"\"{answer}\"\"\"
    """
    return call_ollama_model(prompt, model)


# --- CLEAN AI RESPONSE ---
def clean_ai_response(text):
    unwanted_prefixes = [
        "Sure,", " improved", "improvements made", 
        "revised", "polished", "notes:", "additional"
    ]
    for prefix in unwanted_prefixes:
        text = text.lower().replace(prefix, "")
    return text.strip().capitalize()


# --- FILTER & POLISH QUESTION ---
def filter_and_polish_question(content, model="llama3"):
    prompt = f"""
    You are a helpful assistant.
    
    Determine if the given text is a valid question or irrelevant content.
    Return improved question if valid. If not, return nothing.

    Text:
    \"\"\"{content}\"\"\"
    """
    result = call_ollama_model(prompt, model).strip()

    if not result:
        return None

    irrelevant_keywords = [
        "irrelevant", "not valid", "not a question", "no question", 
        "this is not a question", "not applicable"
    ]

    content_lower = result.lower()
    if any(keyword in content_lower for keyword in irrelevant_keywords):
        return None

    return result


# --- SAVE OPEN QUESTIONS ---
def ai_filtered_open_questions(open_question_list):
    open_doc_path = 'Open_Questions.docx'
    open_doc = Document()
    open_doc.add_heading('Open Questions: ', level=0)

    for question in open_question_list:
        try:
            polished_q = clean_ai_response(filter_and_polish_question(question))
            if polished_q:
                open_doc.add_heading(f"Q: {polished_q}", level=2)
        except Exception:
            open_doc.add_heading(f"Q: {question}", level=2)
            open_doc.add_paragraph("A: Error determining relevance.")
            open_doc.add_paragraph()

    open_doc.save(open_doc_path)
    return open_doc_path


# --- FEEDBACK EXTRACTION ---
def extract_feedback(feedback_files, selected_title):
    combined_df = pd.DataFrame()

    for file in feedback_files:
        df = pd.read_excel(file)
        filtered_df = df[df['Training Title'] == selected_title]
        combined_df = pd.concat([combined_df, filtered_df], ignore_index=True)

    if combined_df.empty:
        raise ValueError(f"No feedback found for the selected title: {selected_title}")

    avg_rating = round(combined_df["User's Course Rating"].mean(), 2)

    feedback_blocks = []
    for _, row in combined_df.iterrows():
        title = str(row.get('User Course Review Title', '')).strip()
        desc = str(row.get('User Course Review Desc', '')).strip()
        if title or desc:
            feedback_blocks.append(f"Title: {title}\nDescription: {desc}")

    combined_feedback = "\n\n".join(feedback_blocks)
    return avg_rating, combined_feedback


# --- AI SUMMARY GENERATOR ---
def generate_ai_summary(session_title, avg_rating, feedback_text, model="llama3"):
    prompt = f"""
    Generate a formal, professional summary of the session based on feedback.
    Avoid email formatting. Keep it concise.

    Feedback:
    {feedback_text}
    """
    return call_ollama_model(prompt, model)


# --- EMAIL BODY ---
def create_email_body(session_title, session_date, attendance_count, avg_rating, ai_summary, repository_link, sharepoint_link, open_questions):
    return f"""
Hi All,

Thank you for your participation in the {session_title} session on {session_date}. The session was highly appreciated, receiving an average star rating of {avg_rating} out of 5.

I am confident that the upcoming sessions will be equally impactful, building on the strong foundation laid by this session.

The session had a strong attendance of {attendance_count}, and here are some highlights from the feedback:
{ai_summary}

We have also compiled a Q&A report from the session, which you can access here: {repository_link}, as well as in this email. There are {open_questions} open questions that require further clarification. Once the Q&A document is finalized, it will be uploaded to the session repository link.

Links updated on SharePoint page {sharepoint_link} for session recording and supporting materials.

Thank you again for your valuable contribution to this training program. Looking forward to the next session in the series.

Best Regards
"""


# --- CREATE EMAIL MSG ---
def create_eml_with_attachments(subject, body, attachment_paths, output_path='Session_Email.eml'):
    msg = EmailMessage()
    msg['Subject'] = subject
    msg['From'] = "you@example.com"
    msg['To'] = "recipient@example.com"
    msg.set_content(body)

    for file_path in attachment_paths:
        mime_type, _ = mimetypes.guess_type(file_path)
        mime_type = mime_type or 'application/octet-stream'
        maintype, subtype = mime_type.split('/', 1)

        with open(file_path, 'rb') as f:
            file_data = f.read()
            file_name = os.path.basename(file_path)

        msg.add_attachment(file_data, maintype=maintype, subtype=subtype, filename=file_name)

    with open(output_path, 'wb') as f:
        f.write(bytes(msg))

    return output_path

# --- Q&A REPORT GENERATOR ---
def generate_qa_report_with_transcript(qna_files, transcript_files=None, model="llama3"):
    dfs = []
    for file in qna_files:
        try:
            dfs.append(pd.read_csv(file))
        except Exception:
            pass
    if not dfs:
        raise ValueError("No valid Q&A files found.")

    df = pd.concat(dfs, ignore_index=True)
    df['Content'] = df['Content'].fillna('')
    df = df.drop_duplicates(subset=['Conversation Id', 'Content'])

    questions = df[df['Type'].isin(['Question', 'Announcement'])]
    responses = df[df['Type'] == 'Response']
    response_dict = responses.groupby('Conversation Id')['Content'].apply(list).to_dict()

    def normalize(text):
        return text.strip().lower().replace('\n', ' ').replace('\r', ' ')

    answered_contents = set(normalize(c) for c in responses['Content'])

    open_question_list = []
    for _, row in questions.iterrows():
        question_text = normalize(row['Content'])
        if not any(question_text in resp for resp in answered_contents):
            open_question_list.append(row['Content'])

    filtered_open_questions = []
    for question in open_question_list:
        polished_question = filter_and_polish_question(question, model)
        if polished_question:
            filtered_open_questions.append(polished_question)

    open_questions = len(filtered_open_questions)
    open_doc_path = ai_filtered_open_questions(filtered_open_questions) if open_questions > 0 else None

    qa_pairs = []
    for _, row in questions.iterrows():
        conv_id = row['Conversation Id']
        question_text = row['Content'].strip()
        answer_list = response_dict.get(conv_id, [])
        structured_answer = '\n'.join([f"- {a.strip()}" for a in answer_list if a.strip()])
        if question_text and structured_answer:
            qa_pairs.append((question_text, structured_answer))

    transcript = parse_vtt_files(transcript_files) if transcript_files else ""

    polished_qa_pairs = []
    for question, raw_answers in qa_pairs:
        raw_lines = raw_answers.lower().split('\n')
        mentioned_as_verbal = any("answered verbally" in line for line in raw_lines)

        full_answer = f"Moderator's Answers:\n{raw_answers}"
        if mentioned_as_verbal:
            transcript_answer = extract_answer_from_transcript(question, transcript)
            full_answer += f"\n\nTranscript-based Answer:\n{transcript_answer}"

        polished_q = clean_ai_response(polish_question(question))
        polished_a = clean_ai_response(polish_answer(polished_q, full_answer))
        polished_qa_pairs.append((polished_q, polished_a))

    qa_doc_path = 'Final_Q&A.docx'
    doc = Document()
    doc.add_heading('Q&A (Including Transcript Analysis)', level=0)
    for question, answer in polished_qa_pairs:
        doc.add_heading(f"Q: {question}", level=2)
        p = doc.add_paragraph(f"A: {answer}")
        p.style.font.size = Pt(12)
        doc.add_paragraph()
    doc.save(qa_doc_path)

    return qa_doc_path, open_doc_path, open_questions


# --- MAIN WORKFLOW ---
def process_training_files(
    attendance_paths,
    qna_paths,
    transcript_paths,
    feedback_paths,
    training_title,
    repository_link,
    sharepoint_link,
    output_dir,
    session_id
):
    feedback_df = validate_training_title(feedback_paths, training_title)
    session_date, attendance_count = extract_attendance_details(attendance_paths)
    # avg_rating, feedback_text = extract_feedback([feedback_df], training_title)
    avg_rating = round(feedback_df["User's Course Rating"].mean(), 2)

    feedback_blocks = []
    for _, row in feedback_df.iterrows():
        title = str(row.get('User Course Review Title', '')).strip()
        desc = str(row.get('User Course Review Desc', '')).strip()
        if title or desc:
            feedback_blocks.append(f"Title: {title}\nDescription: {desc}")
    
    feedback_text = "\n\n".join(feedback_blocks)

    qa_docx, open_docx, open_questions = generate_qa_report_with_transcript(qna_paths, transcript_paths)
    ai_summary = generate_ai_summary(training_title, avg_rating, feedback_text)

    email_body = create_email_body(
        training_title,
        session_date,
        attendance_count,
        avg_rating,
        ai_summary,
        repository_link,
        sharepoint_link,
        open_questions
    )

    qa_output_path = os.path.join(output_dir, f"{session_id}_QnA.docx")
    open_output_path = os.path.join(output_dir, f"{session_id}_Open.docx")
    eml_output_path = os.path.join(output_dir, f"{session_id}_Email.eml")

    os.rename(qa_docx, qa_output_path)
    os.rename(open_docx, open_output_path)

    create_eml_with_attachments(
        subject=f"Summary: {training_title} on {session_date}",
        body=email_body,
        attachment_paths=[qa_output_path, open_output_path],
        output_path=eml_output_path
    )


    return {
        "qa_docx": qa_output_path,
        "open_docx": open_output_path,
        "eml_path": eml_output_path
    }
